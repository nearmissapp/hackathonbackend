# gpt_caller.py

import base64
import os
from openai import OpenAI
from dotenv import load_dotenv
from prompt_generator import AnalyzeImageRisks, FormatRiskAsJson, RetrieveInformation
from docx import Document
from PIL import Image
import io

# .env 파일 로드
load_dotenv()

class GptCaller:
    def __init__(self):
        """
        클래스 초기화 메서드
        """
        self.api_key = os.getenv("OPENAI_API_KEY")
        if not self.api_key:
            raise ValueError("OPENAI_API_KEY가 .env 파일에 설정되지 않았습니다.")

        # OpenRouter API
        self.client = OpenAI(
            # base_url="https://openrouter.ai/api/v1", #다시 openai api로 변경됨.
            api_key=self.api_key
        )

    @staticmethod
    def encode_image_to_base64(image_path, max_size=1024):
        """
        이미지를 리사이즈하고 Base64로 인코딩하는 메서드.
        :param image_path: 인코딩할 이미지 파일 경로
        :param max_size: 이미지의 최대 크기 (픽셀 단위, 기본값: 1024)
        :return: Base64 인코딩된 이미지 문자열
        """
        with open(image_path, "rb") as image_file:
            image = Image.open(image_file)
            width, height = image.size

            if width > max_size or height > max_size:
                if width > height:
                    new_width = max_size
                    new_height = int((max_size / width) * height)
                else:
                    new_height = max_size
                    new_width = int((max_size / height) * width)
                image = image.resize((new_width, new_height), Image.LANCZOS)
                print(f"조정된 이미지 크기: {new_width} x {new_height} 픽셀")  # 조정된 이미지의 픽셀 크기 출력

            # RGBA 모드를 RGB 모드로 변환
            if image.mode == 'RGBA':
                image = image.convert('RGB')

            # 이미지 파일을 메모리에 저장하고 Base64로 인코딩
            with io.BytesIO() as output:
                image.save(output, format="JPEG")
                return base64.b64encode(output.getvalue()).decode("utf-8")

    def analyze_image_risks(self, image_path, comment, location):
        """
        이미지의 잠재적 위험 요소를 분석하는 메서드.
        :param image_path: 분석할 이미지 파일 경로
        :return: OpenAI API 응답
        """
        prompt_manager = AnalyzeImageRisks()
        system_prompt = prompt_manager.get_system_prompt()
        user_prompt = prompt_manager.get_user_prompt(comment, location)
        image_base64 = self.encode_image_to_base64(image_path)

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-2024-11-20",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": user_prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/jpeg;base64,{image_base64}"
                                }
                            }
                        ]
                    }
                ],
                temperature=0.3,
                max_tokens=16000,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0
            )
            # print(response.json())
            return response, image_base64
        except Exception as e:
            return {"error": str(e)}, image_base64

    def format_risk_as_json(self, analyzed_image_risks):
        """
        잠재적 위험 요소 분석 결과를 JSON 형식으로 변환하는 메서드.
        :param analyzed_image_risks: 잠재적 위험 요소 평문 레포트
        :return: JSON 포맷의 OpenAI API 응답
        """
        prompt_manager = FormatRiskAsJson()
        system_prompt = prompt_manager.get_system_prompt()
        user_prompt = prompt_manager.get_user_prompt(analyzed_image_risks)
        tools = [prompt_manager.get_tool()]

        try:
            # 메시지 정의
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ]

            # OpenRouter API 호출
            response = self.client.chat.completions.create(
                model="gpt-4o-2024-11-20",
                messages=messages,
                tools=tools,  # tools 필드 추가
                tool_choice=tools[0],  # 도구 호출 강제
                temperature=0.3,
                max_tokens=16000,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0
            )

            # 결과 반환
            # print(response.json())
            return response
        except Exception as e:
            return {"error": str(e)}

    def retrieve_information(self, formatted_risks_json, doc_search_keyword):
        """
        JSON으로 구성된 잠재적 위험 요소 분석 결과에서 관련 담당자 정보를 검색하고 출력하는 메서드.
        :param formatted_risks_json: JSON 형식의 위험 요소 분석 결과
        :param doc_search_keyword: 검색할 문서 키워드
        :return: OpenAI API 응답
        """

        def read_docx(file_path):
            try:
                # 파일 불러오기
                doc = Document(file_path)
                full_text = []

                # 첫 번째 테이블을 마크다운으로 변환
                if doc.tables:
                    first_table = doc.tables[0]
                    table_md = []

                    # 테이블의 첫 번째 행을 헤더로 사용
                    header_cells = first_table.rows[0].cells
                    header = '| ' + ' | '.join(cell.text for cell in header_cells) + ' |'
                    separator = '| ' + ' | '.join('---' for _ in header_cells) + ' |'
                    table_md.append(header)
                    table_md.append(separator)

                    # 나머지 행을 본문으로 사용
                    for row in first_table.rows[1:]:
                        row_data = '| ' + ' | '.join(cell.text for cell in row.cells) + ' |'
                        table_md.append(row_data)

                    full_text.append('\n'.join(table_md))

                # 모든 문단(paragraphs) 읽기
                for para in doc.paragraphs:
                    full_text.append(para.text)

                return '\n'.join(full_text)
            except Exception as e:
                return f"오류 발생: {e}"

        try:
            doc_text = read_docx(f"documents/잠재위험 사례집 _ {doc_search_keyword}.docx")
        except Exception as e:
            print(f"문서 읽기 중 오류 발생: {str(e)}")
            return {"error": f"문서 읽기 중 오류 발생: {str(e)}"}

        prompt_manager = RetrieveInformation()
        system_prompt = prompt_manager.get_system_prompt()
        user_prompt = prompt_manager.get_user_prompt(doc_search_keyword, doc_text, formatted_risks_json)
        tools = [prompt_manager.get_tool()]

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-2024-11-20",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                tools=tools,
                tool_choice=tools[0],  # 도구 호출 강제
                temperature=0.3,
                max_tokens=16000,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0
            )
            # print(response.json())
            return response
        except Exception as e:
            print(e)
            return {"error": str(e)}